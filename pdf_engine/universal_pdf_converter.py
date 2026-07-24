"""
Universal PDF → DOCX Converter — 自适应路由 + 双引擎 + 中间态渲染

架构:
  Router (fitz 文本密度分析)
    ├─ STANDARD  → pdf2docx (毫秒级，矢量保真)
    ├─ SCANNED   → docling (视觉模型, OCR + 布局)
    ├─ COMPLEX   → docling (多栏/公式/表格识别)
    └─ FALLBACK  → 逐页截图嵌入 (终极兜底)

中间态: Markdown + JSON metadata (图片位置/表格结构)
渲染层: python-docx 自定义模板

依赖: 纯本地离线，零 API 调用，零数据上云
"""

import os
import sys
import json
import logging
import tempfile
import re
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional, List, Dict, Tuple
from concurrent.futures import ProcessPoolExecutor, as_completed

import fitz  # PyMuPDF
from pdf2docx import Converter as Pdf2DocxConverter

logging.basicConfig(level=logging.INFO, format="[%(levelname)s] %(message)s")
logger = logging.getLogger("UniversalPDFConverter")


# ============================================================
# Data Models
# ============================================================

@dataclass
class PageAnalysis:
    """单页 PDF 分析结果"""
    page_num: int
    char_count: int
    image_count: int
    has_table: bool
    has_formula: bool
    has_multi_column: bool
    text_density: float  # 文字覆盖率 0~1

@dataclass
class PDFInspectionResult:
    """PDF 整体检测结果"""
    total_pages: int
    total_chars: int
    avg_chars_per_page: float
    pages: List[PageAnalysis] = field(default_factory=list)
    classification: str = "STANDARD"  # STANDARD | SCANNED | COMPLEX
    metadata: Dict = field(default_factory=dict)


# ============================================================
# Router — PDF 分析 & 分类
# ============================================================

class PDFInspector:
    """利用 PyMuPDF 进行文本密度与结构特征分析"""

    # 判定阈值
    MIN_CHARS_FOR_TEXT_PAGE = 50       # 单页少于 50 字视为图片页
    SCANNED_RATIO_THRESHOLD = 0.3      # 文字页占比低于 30% 视为扫描件
    MULTI_COLUMN_LINE_THRESHOLD = 3    # 检测到 ≥3 个独立文本列视为多栏

    def __init__(self, pdf_path: str):
        self.pdf_path = pdf_path
        self.doc: Optional[fitz.Document] = None

    def inspect(self) -> PDFInspectionResult:
        """主入口: 分析 PDF 并返回分类结果"""
        self.doc = fitz.open(self.pdf_path)
        total_pages = self.doc.page_count
        pages: List[PageAnalysis] = []
        total_chars = 0
        scanned_page_count = 0
        complex_page_count = 0

        for i in range(total_pages):
            page = self.doc[i]
            analysis = self._analyze_page(page, i + 1)
            pages.append(analysis)
            total_chars += analysis.char_count

            if analysis.char_count < self.MIN_CHARS_FOR_TEXT_PAGE:
                scanned_page_count += 1
            if analysis.has_multi_column or analysis.has_formula or analysis.has_table:
                complex_page_count += 1

        self.doc.close()

        avg_chars = total_chars / max(total_pages, 1)
        scanned_ratio = scanned_page_count / max(total_pages, 1)

        # 分类逻辑
        if scanned_ratio >= self.SCANNED_RATIO_THRESHOLD:
            classification = "SCANNED"
        elif complex_page_count >= total_pages * 0.2:
            classification = "COMPLEX"
        else:
            classification = "STANDARD"

        result = PDFInspectionResult(
            total_pages=total_pages,
            total_chars=total_chars,
            avg_chars_per_page=avg_chars,
            pages=pages,
            classification=classification,
        )

        logger.info(
            f"PDF 分析完成: {total_pages} 页, "
            f"平均 {avg_chars:.0f} 字/页, "
            f"分类 → {classification}"
        )
        return result

    def _analyze_page(self, page: fitz.Page, page_num: int) -> PageAnalysis:
        """分析单页特征"""
        # 文字量
        text = page.get_text()
        char_count = len(text.strip())

        # 图片检测
        image_list = page.get_images(full=True)
        image_count = len(image_list)

        # 文本块分布 (用于多栏检测)
        blocks = page.get_text("blocks")
        text_blocks = [b for b in blocks if b[6] == 0]  # type 0 = text

        # 简易多栏检测: 检查同一 y 坐标范围内是否有多个文本块
        has_multi_column = self._detect_multi_column(text_blocks)

        # 公式/表格检测 (简易启发式)
        has_table = self._detect_table(text)
        has_formula = self._detect_formula(text)

        # 文字覆盖率
        page_rect = page.rect
        page_area = page_rect.width * page_rect.height
        text_area = sum((b[2] - b[0]) * (b[3] - b[1]) for b in text_blocks)
        text_density = min(text_area / max(page_area, 1), 1.0)

        return PageAnalysis(
            page_num=page_num,
            char_count=char_count,
            image_count=image_count,
            has_table=has_table,
            has_formula=has_formula,
            has_multi_column=has_multi_column,
            text_density=text_density,
        )

    def _detect_multi_column(self, text_blocks: List) -> bool:
        """检测多栏排版: 同一水平带内 ≥3 个独立文本块"""
        if len(text_blocks) < self.MULTI_COLUMN_LINE_THRESHOLD:
            return False

        # 按 y 坐标排序
        sorted_blocks = sorted(text_blocks, key=lambda b: b[1])

        # 扫描同一水平带
        band_height = 5  # y 坐标容差
        for i, block in enumerate(sorted_blocks):
            y0, y1 = block[1], block[3]
            same_band = [block]
            for other in sorted_blocks[i + 1:]:
                if abs(other[1] - y0) < band_height or abs(other[3] - y1) < band_height:
                    same_band.append(other)
            if len(same_band) >= self.MULTI_COLUMN_LINE_THRESHOLD:
                return True
        return False

    def _detect_table(self, text: str) -> bool:
        """简易表格检测: 检查是否有规整的表格线或制表符"""
        # 检查制表符
        tab_lines = [l for l in text.split('\n') if l.count('\t') >= 2]
        if len(tab_lines) >= 2:
            return True
        # 检查竖线分隔符
        pipe_lines = [l for l in text.split('\n') if l.count('|') >= 3]
        if len(pipe_lines) >= 2:
            return True
        return False

    def _detect_formula(self, text: str) -> bool:
        """简易公式检测: LaTeX 标记或数学符号密集"""
        latex_indicators = [r'\frac', r'\sum', r'\int', r'\sqrt', r'\alpha',
                            r'\beta', r'\theta', r'\pi', r'\Delta', r'\partial']
        count = sum(text.count(ind) for ind in latex_indicators)
        return count >= 3


# ============================================================
# Engine 1 — STANDARD: pdf2docx 轻量规则引擎
# ============================================================

class StandardEngine:
    """
    适用于: 原生电子 PDF，排版简单，文字可直接提取
    特点: 毫秒级响应，保留矢量字体信息
    """

    def convert(self, pdf_path: str, output_path: str) -> bool:
        logger.info(f"[STANDARD] 使用 pdf2docx 转换: {pdf_path}")
        try:
            cv = Pdf2DocxConverter(pdf_path)
            cv.convert(output_path, start=0, end=None)
            cv.close()
            logger.info(f"[STANDARD] 转换完成 → {output_path}")
            return True
        except Exception as e:
            logger.error(f"[STANDARD] 转换失败: {e}")
            return False


# ============================================================
# Engine 2 — SCANNED/COMPLEX: Docling 视觉引擎
# ============================================================

class VisionEngine:
    """
    适用于: 扫描件、纯图片 PDF、多栏排版、含公式表格
    特点: 本地视觉模型，OCR + 布局分析 + Table/LaTeX 识别

    Docling 输出 Markdown，其中:
      - 表格 → HTML <table> 嵌入 Markdown
      - 公式 → LaTeX $$...$$ 或 $...$
      - 图片 → base64 或引用路径
    """

    def __init__(self):
        self._converter = None

    @property
    def converter(self):
        """延迟加载 Docling (首次加载需下载模型 ~2GB)"""
        if self._converter is None:
            logger.info("[VISION] 加载 Docling 模型 (首次运行需下载 ~2GB)...")
            from docling.document_converter import DocumentConverter
            self._converter = DocumentConverter()
        return self._converter

    def convert(self, pdf_path: str, output_path: str) -> bool:
        logger.info(f"[VISION] 使用 Docling 解析: {pdf_path}")
        try:
            result = self.converter.convert(pdf_path)
            markdown_text = result.document.export_to_markdown()

            # 保存中间态 Markdown (调试用)
            md_path = output_path.replace('.docx', '_intermediate.md')
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(markdown_text)
            logger.info(f"[VISION] Markdown 中间态 → {md_path}")

            # 尝试从原 PDF 提取正文字体名
            base_font = self._detect_dominant_font(pdf_path)

            # Markdown + metadata → DOCX
            success = self._render_docx(markdown_text, output_path, base_font=base_font)
            return success
        except Exception as e:
            logger.error(f"[VISION] 转换失败: {e}")
            return False

    def _detect_dominant_font(self, pdf_path: str) -> Optional[str]:
        """从 PDF 第一页检测主要字体，映射到系统可用字体"""
        try:
            doc = fitz.open(pdf_path)
            page = doc[0]
            blocks = page.get_text("dict")["blocks"]

            font_counts: Dict[str, int] = {}
            for block in blocks:
                if "lines" in block:
                    for line in block["lines"]:
                        for span in line["spans"]:
                            fname = span.get("font", "")
                            if fname:
                                font_counts[fname] = font_counts.get(fname, 0) + len(span["text"])

            doc.close()

            if font_counts:
                dominant = max(font_counts, key=font_counts.get)
                mapped = self._map_font(dominant)
                if mapped:
                    logger.info(f"[VISION] 检测到主字体: {dominant} → {mapped}")
                    return mapped
                logger.info(f"[VISION] 检测到主字体: {dominant} (无需映射)")
                # 清理 PDF 子集化前缀
                clean = dominant.split('+')[-1]
                return clean if clean != dominant else None
        except Exception as e:
            logger.warning(f"[VISION] 字体检测失败: {e}")
        return None

    # PDF 常见字体 → Word 可用字体映射
    FONT_MAP = {
        # 中文字体
        'SimSun': '宋体', 'SimHei': '黑体', 'KaiTi': '楷体',
        'FangSong': '仿宋', 'Microsoft YaHei': '微软雅黑',
        'PMingLiU': '新細明體', 'MingLiU': '細明體',
        # 英文字体
        'TimesNewRomanPSMT': 'Times New Roman', 'TimesNewRoman': 'Times New Roman',
        'ArialMT': 'Arial', 'Arial': 'Arial',
        'Helvetica': 'Arial', 'Helvetica-Bold': 'Arial',
        'CourierNew': 'Courier New', 'Courier': 'Courier New',
        # 日韩
        'MS-Mincho': 'MS Mincho', 'MS-Gothic': 'MS Gothic',
        'Gulim': 'Gulim', 'Batang': 'Batang',
    }

    @classmethod
    def _map_font(cls, pdf_font_name: str) -> Optional[str]:
        """将 PDF 内嵌字体名映射到系统可用字体"""
        if not pdf_font_name:
            return None
        # 去掉 PDF 子集化前缀 (如 ABCDEF+TimesNewRoman → TimesNewRoman)
        clean = pdf_font_name.split('+')[-1].replace('-', '')
        return cls.FONT_MAP.get(clean) or cls.FONT_MAP.get(pdf_font_name)

    def _render_docx(self, markdown_text: str, output_path: str,
                     base_font: str = None) -> bool:
        """
        将 Markdown 渲染为 DOCX

        Args:
            markdown_text: Docling 输出的 Markdown
            output_path: 输出 .docx 路径
            base_font: 正文字体 (不指定则用 Word 默认主题字体)
        """
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn

        doc = Document()

        # 默认正文样式：不硬编码字体名，让 Word 用主题默认字体
        # 仅设置字号，确保不同语言环境下字体显示正常
        style = doc.styles['Normal']
        font = style.font
        font.size = Pt(11)
        if base_font:
            font.name = base_font

        # 分段解析 Markdown
        lines = markdown_text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]

            # 跳过空行
            if not line.strip():
                i += 1
                continue

            # HTML 表格 (Docling 用 HTML 嵌表格)
            if line.strip().startswith('<table>'):
                table_html, i = self._extract_html_table(lines, i)
                self._add_table_from_html(doc, table_html)
                continue

            # LaTeX 公式块 $$...$$
            if line.strip().startswith('$$'):
                formula_lines, i = self._extract_block(lines, i, '$$', '$$')
                formula_text = '\n'.join(formula_lines)
                self._add_formula_paragraph(doc, formula_text, block=True)
                continue

            # 行内公式 $...$
            if '$' in line:
                self._add_formatted_paragraph(doc, line)
                i += 1
                continue

            # 标题
            if line.startswith('### '):
                self._add_heading(doc, line[4:], level=3)
            elif line.startswith('## '):
                self._add_heading(doc, line[3:], level=2)
            elif line.startswith('# '):
                self._add_heading(doc, line[2:], level=1)
            else:
                # 普通段落
                self._add_formatted_paragraph(doc, line)

            i += 1

        doc.save(output_path)
        logger.info(f"[VISION] DOCX 渲染完成 → {output_path}")
        return True

    def _extract_html_table(self, lines: List[str], start: int) -> Tuple[str, int]:
        """提取完整的 HTML 表格块"""
        html = lines[start]
        i = start + 1
        while i < len(lines) and '</table>' not in html:
            html += '\n' + lines[i]
            i += 1
        return html, i

    def _extract_block(self, lines: List[str], start: int,
                       _open_marker: str, close_marker: str) -> Tuple[List[str], int]:
        """提取标记块"""
        block_lines = []
        i = start + 1
        while i < len(lines) and close_marker not in lines[i]:
            block_lines.append(lines[i])
            i += 1
        return block_lines, i + 1

    def _add_heading(self, doc, text: str, level: int):
        """添加标题"""
        from docx.shared import Pt
        heading = doc.add_heading(text, level=level)
        return heading

    def _add_formula_paragraph(self, doc, formula: str, block: bool = False):
        """添加公式段落 (保留 LaTeX 原文，Word 中用 Cambria Math 渲染)"""
        from docx.shared import Pt
        from docx.oxml.ns import qn
        p = doc.add_paragraph()
        p.alignment = 1  # 居中
        run = p.add_run(formula.strip())
        run.font.name = 'Cambria Math'
        run.font.size = Pt(11)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Cambria Math')
        return p

    def _add_formatted_paragraph(self, doc, text: str):
        """添加带格式的段落 (处理行内 Markdown 标记)"""
        from docx.shared import Pt
        p = doc.add_paragraph()

        # 简易 Markdown 行内解析: **粗体**, *斜体*, `代码`
        segments = re.split(r'(\*\*.*?\*\*|\*[^*].*?[^*]\*|`.*?`|\$.*?\$)', text)
        for seg in segments:
            if seg.startswith('**') and seg.endswith('**'):
                run = p.add_run(seg[2:-2])
                run.bold = True
            elif seg.startswith('*') and seg.endswith('*') and not seg.startswith('**'):
                run = p.add_run(seg[1:-1])
                run.italic = True
            elif seg.startswith('`') and seg.endswith('`'):
                run = p.add_run(seg[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            elif seg.startswith('$') and seg.endswith('$'):
                run = p.add_run(seg[1:-1])
                run.font.name = 'Cambria Math'
                run.font.size = Pt(11)
            else:
                p.add_run(seg)
        return p

    def _add_table_from_html(self, doc, html: str):
        """从 HTML 表格创建 python-docx Table"""
        from docx.shared import Pt, Inches
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, 'html.parser')
        table_elem = soup.find('table')
        if not table_elem:
            return

        rows = table_elem.find_all('tr')
        if not rows:
            return

        # 计算最大列数
        max_cols = 0
        for row in rows:
            cols = row.find_all(['td', 'th'])
            max_cols = max(max_cols, len(cols))

        if max_cols == 0:
            return

        table = doc.add_table(rows=len(rows), cols=max_cols, style='Table Grid')

        for r_idx, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for c_idx, cell in enumerate(cells):
                if c_idx < max_cols:
                    table.cell(r_idx, c_idx).text = cell.get_text(strip=True)

        doc.add_paragraph()  # 表格后空行


# ============================================================
# Engine 3 — FALLBACK: 逐页截图嵌入 (终极兜底)
# ============================================================

class FallbackEngine:
    """
    适用于: 所有引擎都失败时的最终兜底
    特点: PyMuPDF 逐页渲染为 PNG，嵌入 DOCX
         视觉完全保真，但文字不可编辑
    """

    def convert(self, pdf_path: str, output_path: str) -> bool:
        logger.info(f"[FALLBACK] 使用逐页截图嵌入: {pdf_path}")
        try:
            from docx import Document
            from docx.shared import Inches

            doc = Document()
            pdf = fitz.open(pdf_path)

            # A4 宽度 6.5 英寸 (减去边距)
            page_width = Inches(6.5)

            for i in range(pdf.page_count):
                page = pdf[i]
                # 渲染为 PNG, DPI=150 平衡质量与大小
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("png")

                # 计算图片高度保持纵横比
                aspect = pix.height / pix.width
                page_height = Inches(6.5 * aspect)

                # 写入临时 PNG 文件
                tmp_png = tempfile.mkstemp(suffix='.png')[1]
                with open(tmp_png, 'wb') as f:
                    f.write(img_bytes)

                # 嵌入 DOCX
                para = doc.add_paragraph()
                para.alignment = 1  # 居中
                run = para.add_run()
                run.add_picture(tmp_png, width=page_width, height=page_height)

                # 清理临时文件
                os.unlink(tmp_png)

                logger.info(f"[FALLBACK] 第 {i+1}/{pdf.page_count} 页完成")

            pdf.close()
            doc.save(output_path)
            logger.info(f"[FALLBACK] 转换完成 → {output_path}")
            return True
        except Exception as e:
            logger.error(f"[FALLBACK] 转换失败: {e}")
            return False


# ============================================================
# 主服务 — UniversalPDFConverter
# ============================================================

class UniversalPDFConverter:
    """
    PDF → DOCX 统一转换入口

    路由策略:
      1. inspect_pdf() 分析 PDF，分类为 STANDARD / SCANNED / COMPLEX
      2. STANDARD  → StandardEngine (pdf2docx)
      3. SCANNED   → VisionEngine (docling)
      4. COMPLEX   → VisionEngine (docling)
      5. 任何引擎失败 → FallbackEngine (逐页截图)
    """

    def __init__(self, cache_dir: Optional[str] = None):
        """
        Args:
            cache_dir: 模型缓存目录 (docling 模型下载位置)
                       默认 ~/.cache/universal_pdf_converter
        """
        self.cache_dir = cache_dir or os.path.expanduser(
            "~/.cache/universal_pdf_converter"
        )
        os.makedirs(self.cache_dir, exist_ok=True)

        self.inspector: Optional[PDFInspector] = None
        self._vision_engine: Optional[VisionEngine] = None
        self._standard_engine: Optional[StandardEngine] = None
        self._fallback_engine: Optional[FallbackEngine] = None

    # ---- 懒加载引擎 (减少内存占用) ----

    @property
    def standard_engine(self) -> StandardEngine:
        if self._standard_engine is None:
            self._standard_engine = StandardEngine()
        return self._standard_engine

    @property
    def vision_engine(self) -> VisionEngine:
        if self._vision_engine is None:
            self._vision_engine = VisionEngine()
        return self._vision_engine

    @property
    def fallback_engine(self) -> FallbackEngine:
        if self._fallback_engine is None:
            self._fallback_engine = FallbackEngine()
        return self._fallback_engine

    def inspect_pdf(self, pdf_path: str) -> PDFInspectionResult:
        """分析 PDF 并返回路由分类"""
        self.inspector = PDFInspector(pdf_path)
        return self.inspector.inspect()

    def convert(self, pdf_path: str, output_path: Optional[str] = None) -> str:
        """
        主入口: 自适应路由转换 PDF → DOCX

        Args:
            pdf_path: 输入 PDF 路径
            output_path: 输出 DOCX 路径 (可选，默认同名 .docx)

        Returns:
            输出 DOCX 文件路径
        """
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF 不存在: {pdf_path}")

        if output_path is None:
            output_path = pdf_path.rsplit('.', 1)[0] + '.docx'

        # Step 1: 分析 & 路由
        inspection = self.inspect_pdf(pdf_path)
        classification = inspection.classification
        logger.info(f"路由决策: {classification} → {pdf_path}")

        # Step 2: 按分类调用引擎 + 异常降级
        success = False

        if classification == "STANDARD":
            success = self.standard_engine.convert(pdf_path, output_path)
            if not success:
                logger.warning("[ROUTER] STANDARD 引擎失败，降级到 VISION 引擎")
                success = self.vision_engine.convert(pdf_path, output_path)

        elif classification in ("SCANNED", "COMPLEX"):
            success = self.vision_engine.convert(pdf_path, output_path)
            if not success:
                logger.warning("[ROUTER] VISION 引擎失败，降级到 FALLBACK 引擎")
                success = self.fallback_engine.convert(pdf_path, output_path)

        else:
            # 兜底: 直接尝试 FALLBACK
            success = self.fallback_engine.convert(pdf_path, output_path)

        # Step 3: 最终兜底
        if not success:
            logger.error("[ROUTER] 所有引擎均失败，使用 FALLBACK 引擎")
            success = self.fallback_engine.convert(pdf_path, output_path)

        if not success:
            raise RuntimeError(f"PDF 转换完全失败: {pdf_path}")

        logger.info(f"转换成功: {pdf_path} → {output_path}")
        return output_path

    def convert_batch(self, pdf_paths: List[str], output_dir: str,
                      max_workers: int = 2) -> Dict[str, str]:
        """
        批量转换 (多进程)

        Args:
            pdf_paths: PDF 文件路径列表
            output_dir: 输出目录
            max_workers: 最大并行进程数

        Returns:
            {输入路径: 输出路径} 字典
        """
        os.makedirs(output_dir, exist_ok=True)
        results = {}

        with ProcessPoolExecutor(max_workers=max_workers) as executor:
            futures = {}
            for pdf_path in pdf_paths:
                out_name = Path(pdf_path).stem + '.docx'
                out_path = os.path.join(output_dir, out_name)
                future = executor.submit(self._convert_worker, pdf_path, out_path)
                futures[future] = pdf_path

            for future in as_completed(futures):
                pdf_path = futures[future]
                try:
                    out_path = future.result()
                    results[pdf_path] = out_path
                    logger.info(f"[BATCH] ✓ {pdf_path}")
                except Exception as e:
                    logger.error(f"[BATCH] ✗ {pdf_path}: {e}")
                    results[pdf_path] = None

        return results

    @staticmethod
    def _convert_worker(pdf_path: str, output_path: str) -> str:
        """独立进程中的转换 worker (避免 GIL + 内存隔离)"""
        converter = UniversalPDFConverter()
        return converter.convert(pdf_path, output_path)

    def cleanup(self):
        """释放引擎资源 (长驻服务调用)"""
        if self._vision_engine and self._vision_engine._converter:
            del self._vision_engine._converter
            self._vision_engine._converter = None
        if self._standard_engine:
            self._standard_engine = None
        if self._fallback_engine:
            self._fallback_engine = None
        import gc
        gc.collect()


# ============================================================
# CLI 入口
# ============================================================

def main():
    """命令行入口"""
    import argparse

    parser = argparse.ArgumentParser(
        description="Universal PDF → DOCX Converter",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python universal_pdf_converter.py input.pdf
  python universal_pdf_converter.py input.pdf -o output.docx
  python universal_pdf_converter.py input.pdf --inspect-only
  python universal_pdf_converter.py --batch dir/ --output out_dir/
        """,
    )
    parser.add_argument("input", nargs='?', help="输入 PDF 文件路径")
    parser.add_argument("-o", "--output", help="输出 DOCX 路径")
    parser.add_argument("--inspect-only", action="store_true",
                        help="仅分析 PDF 不转换")
    parser.add_argument("--batch", help="批量模式: 输入目录")
    parser.add_argument("--output-dir", help="批量模式: 输出目录")
    parser.add_argument("--workers", type=int, default=2,
                        help="批量模式: 并行进程数 (默认 2)")
    parser.add_argument("--verbose", "-v", action="store_true",
                        help="详细日志")

    args = parser.parse_args()

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    converter = UniversalPDFConverter()

    if args.inspect_only and args.input:
        result = converter.inspect_pdf(args.input)
        print(json.dumps({
            "classification": result.classification,
            "total_pages": result.total_pages,
            "total_chars": result.total_chars,
            "avg_chars_per_page": round(result.avg_chars_per_page, 1),
            "pages": [
                {
                    "page": p.page_num,
                    "chars": p.char_count,
                    "images": p.image_count,
                    "table": p.has_table,
                    "formula": p.has_formula,
                    "multi_column": p.has_multi_column,
                    "density": round(p.text_density, 3),
                }
                for p in result.pages
            ]
        }, indent=2, ensure_ascii=False))
        return

    if args.batch:
        pdf_files = list(Path(args.batch).glob("*.pdf"))
        if not pdf_files:
            print(f"目录 {args.batch} 中没有找到 PDF 文件")
            sys.exit(1)
        output_dir = args.output_dir or os.path.join(args.batch, "docx_output")
        results = converter.convert_batch(
            [str(f) for f in pdf_files],
            output_dir,
            max_workers=args.workers
        )
        success_count = sum(1 for v in results.values() if v is not None)
        print(f"\n批量转换完成: {success_count}/{len(results)} 成功")
        if success_count != len(results):
            sys.exit(1)
        return

    if args.input:
        output_path = converter.convert(args.input, args.output)
        print(f"[OK] {output_path}")
        return

    parser.print_help()


if __name__ == "__main__":
    main()
